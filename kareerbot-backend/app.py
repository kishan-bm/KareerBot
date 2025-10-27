import os
import json
import re
import shutil
from flask import Flask, request, jsonify
from flask_cors import CORS
from dotenv import load_dotenv
import hashlib
from datetime import datetime, timedelta
import jwt

# LangChain and Gemini Imports (optional)
try:
    from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.docstore.document import Document
    from langchain.chains import create_retrieval_chain
    from langchain.chains.combine_documents.stuff import create_stuff_documents_chain
    from langchain_core.prompts import ChatPromptTemplate
    from langchain_community.vectorstores import Chroma
    from langchain_core.prompts import MessagesPlaceholder
    from langchain_community.tools.tavily_search import TavilySearchResults
    from langchain.agents import AgentExecutor, create_tool_calling_agent
    from langchain_core.prompts import PromptTemplate
    LANGCHAIN_AVAILABLE = True
except Exception as e:
    # Graceful degradation for environments without the LLM libs installed.
    LANGCHAIN_AVAILABLE = False
    # Provide minimal stubs so the app can start for auth and DB work.
    class ChatGoogleGenerativeAI:
        def __init__(*args, **kwargs):
            pass
        def invoke(self, prompt):
            class R: content = '{}';
            return R()

    class GoogleGenerativeAIEmbeddings:
        def __init__(*args, **kwargs):
            pass

    def RecursiveCharacterTextSplitter(*args, **kwargs):
        class Splitter:
            def split_documents(self, docs):
                return docs
        return Splitter()

    class Document:
        def __init__(self, page_content=''):
            self.page_content = page_content

    def create_retrieval_chain(*args, **kwargs):
        return None

    def create_stuff_documents_chain(*args, **kwargs):
        return None

    class Chroma:
        def __init__(*args, **kwargs):
            pass
        def persist(*args, **kwargs):
            pass
        def add_documents(*args, **kwargs):
            pass
        @classmethod
        def from_documents(cls, *args, **kwargs):
            return cls()

    class MessagesPlaceholder:
        def __init__(self, *args, **kwargs):
            pass

    class TavilySearchResults:
        def __init__(self, *args, **kwargs):
            pass

    class AgentExecutor:
        def __init__(self, *args, **kwargs):
            pass

    def create_tool_calling_agent(*args, **kwargs):
        return None

# File parsing imports
import pypdf
from docx import Document as DocxDocument
from sqlalchemy import create_engine, Column, Integer, String, DateTime, text
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
from werkzeug.security import generate_password_hash, check_password_hash


load_dotenv()
app = Flask(__name__)
CORS(app)

vector_store = None
genai_api_key = os.getenv("GEMINI_API_KEY")

# --- Database setup (Postgres) ---
DATABASE_URL = os.getenv('DATABASE_URL') or 'sqlite:///local_dev.db'
JWT_SECRET = os.getenv('JWT_SECRET') or 'dev_jwt_secret'
engine = create_engine(DATABASE_URL, echo=False, future=True)
SessionLocal = sessionmaker(bind=engine)
Base = declarative_base()


class User(Base):
    __tablename__ = 'users'
    id = Column(Integer, primary_key=True, index=True)
    username = Column(String, unique=True, index=True, nullable=True)
    email = Column(String, unique=True, index=True, nullable=True)
    phone = Column(String, unique=True, index=True, nullable=True)
    password_hash = Column(String, nullable=False)
    created_at = Column(DateTime, default=datetime.utcnow)


Base.metadata.create_all(bind=engine)

# Ensure `email` and `phone` columns exist on older DBs created before they were added to the model.
def ensure_user_columns():
    try:
        with engine.begin() as conn:
            # check email
            has_email = conn.execute(text("SELECT column_name FROM information_schema.columns WHERE table_name='users' AND column_name='email'"))
            if not has_email.scalar():
                try:
                    conn.execute(text('ALTER TABLE users ADD COLUMN email VARCHAR'))
                except Exception:
                    pass

            # check phone
            has_phone = conn.execute(text("SELECT column_name FROM information_schema.columns WHERE table_name='users' AND column_name='phone'"))
            if not has_phone.scalar():
                try:
                    conn.execute(text('ALTER TABLE users ADD COLUMN phone VARCHAR'))
                except Exception:
                    pass
    except Exception:
        # If anything goes wrong (e.g., using sqlite), skip migration
        pass

ensure_user_columns()

# Initialize chat model and embeddings
chat_model = ChatGoogleGenerativeAI(model="gemini-1.5-flash", google_api_key=genai_api_key)
embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=genai_api_key)

# Tavily Search Tool Setup
tavily_api_key = os.getenv("TAVILY_API_KEY")
search = TavilySearchResults(api_key=tavily_api_key)
tools = [search]

# Prompt for the tool-calling agent
agent_prompt = ChatPromptTemplate.from_messages(
    [
        ("system", "You are a helpful AI assistant that can use tools."),
        ("human", "{input}"),
        MessagesPlaceholder(variable_name="agent_scratchpad"),
    ]
)

# Create a tool-calling agent
agent = create_tool_calling_agent(chat_model, tools, agent_prompt)
agent_executor = AgentExecutor(agent=agent, tools=tools, verbose=True)

CHROMA_DB_PATH = "./chroma_db"

def get_pdf_text(pdf_file):
    reader = pypdf.PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

def get_docx_text(docx_file):
    doc = DocxDocument(docx_file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text


def _ingested_path_for(user_id: str):
    safe = user_id if user_id else 'default'
    return f"ingested_docs_{safe}.json"


def load_ingested_docs(user_id: str = 'default'):
    path = _ingested_path_for(user_id)
    if not os.path.exists(path):
        return []
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return []


def save_ingested_docs(docs, user_id: str = 'default'):
    path = _ingested_path_for(user_id)
    try:
        with open(path, 'w', encoding='utf-8') as f:
            json.dump(docs, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"Failed to save ingested docs for {user_id}: {e}")


def get_user_id_from_request(req):
    # Prefer Authorization Bearer <token>
    auth = None
    if req.headers and 'Authorization' in req.headers:
        auth = req.headers.get('Authorization')
    elif req.args and 'authorization' in req.args:
        auth = req.args.get('authorization')
    if auth and auth.startswith('Bearer '):
        token = auth.split(' ', 1)[1]
        try:
            payload = jwt.decode(token, JWT_SECRET, algorithms=['HS256'])
            return str(payload.get('user_id'))
        except Exception:
            return None
    return None


@app.route('/api/save-plan', methods=['POST'])
def save_plan():
    data = request.json or {}
    user_id = get_user_id_from_request(request) or data.get('user_id') or request.args.get('user_id') or 'default'
    plan = data.get('plan')
    if not plan:
        return jsonify({'error': 'plan is required'}), 400
    path = f'saved_plan_{user_id}.json'
    try:
        with open(path, 'w', encoding='utf-8') as f:
            json.dump({'plan': plan, 'saved_at': datetime.utcnow().isoformat()}, f, ensure_ascii=False, indent=2)
        return jsonify({'status': 'ok', 'path': path})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/register', methods=['POST'])
def register():
    data = request.json or {}
    # Expecting: { contact: <email or phone>, password: <password>, username?: <display name> }
    contact = data.get('contact')
    password = data.get('password')
    username = data.get('username')
    if not contact or not password:
        return jsonify({'error': 'contact (email or phone) and password required'}), 400
    db = SessionLocal()
    # determine if contact is email or phone
    is_email = '@' in contact
    if is_email:
        existing = db.query(User).filter(User.email == contact).first()
    else:
        existing = db.query(User).filter(User.phone == contact).first()
    if existing:
        return jsonify({'error': 'user already exists'}), 400
    # create user
    user = User(username=username or None, email=contact if is_email else None, phone=contact if not is_email else None, password_hash=generate_password_hash(password))
    db.add(user)
    db.commit()
    db.refresh(user)
    # create a token
    token = jwt.encode({'user_id': user.id, 'exp': datetime.utcnow() + timedelta(days=30)}, JWT_SECRET, algorithm='HS256')
    return jsonify({'status': 'ok', 'user_id': user.id, 'token': token})


@app.route('/api/login', methods=['POST'])
def login():
    data = request.json or {}
    # Expecting: { contact: <email or phone>, password: <password> }
    contact = data.get('contact')
    password = data.get('password')
    if not contact or not password:
        return jsonify({'error': 'contact and password required'}), 400
    db = SessionLocal()
    is_email = '@' in contact
    if is_email:
        user = db.query(User).filter(User.email == contact).first()
    else:
        user = db.query(User).filter(User.phone == contact).first()
    if not user:
        return jsonify({'error': 'user does not exist'}), 404
    if not check_password_hash(user.password_hash, password):
        return jsonify({'error': 'invalid credentials'}), 401
    token = jwt.encode({'user_id': user.id, 'exp': datetime.utcnow() + timedelta(days=30)}, JWT_SECRET, algorithm='HS256')
    return jsonify({'status': 'ok', 'user_id': user.id, 'token': token})


@app.route('/api/load-plan', methods=['GET'])
def load_plan():
    user_id = get_user_id_from_request(request) or request.args.get('user_id') or 'default'
    path = f'saved_plan_{user_id}.json'
    if not os.path.exists(path):
        return jsonify({'plan': None})
    try:
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return jsonify(data)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/compare-profile', methods=['GET'])
def compare_profile():
    user_id = get_user_id_from_request(request) or request.args.get('user_id') or 'default'
    ingested = load_ingested_docs(user_id)
    if not ingested:
        return jsonify({'error': 'no ingested docs for user', 'suggestion': 'Upload resume or skills first'}), 404

    # Combine the texts into a single prompt for analysis
    combined = '\n\n'.join([d.get('text', '') for d in ingested])
    prompt = f"""
You are a career analyst. Given the user's combined profile and captured skills below, produce a JSON object with:
- summary: one-paragraph summary
- gaps: array of {{title, description, keywords}}
- recommended_next_steps: 10-12 granular micro-steps (title, description, keywords, actions)

User data:
{combined}
"""

    try:
        ai_resp = chat_model.invoke(prompt).content
        # try to extract JSON
        jm = re.search(r'\{.*\}', ai_resp, re.DOTALL)
        if jm:
            analysis = json.loads(jm.group(0))
        else:
            analysis = {'raw': ai_resp}
    except Exception as e:
        analysis = {'error': str(e)}

    return jsonify({'analysis': analysis, 'ingested_count': len(ingested)})

@app.route("/api/process-resume", methods=["POST"])
def process_resume():
    global vector_store

    # Append mode: if vector_store exists, we will add new documents to it.
    # Load existing ingested_docs to append this new source as well.
    # determine user id from form/json/args
    user_id = get_user_id_from_request(request) or (
        request.form.get('user_id') if request.form and 'user_id' in request.form else (
            request.json.get('user_id') if request.json and 'user_id' in request.json else request.args.get('user_id', 'default')
        )
    )

    ingested_docs = load_ingested_docs(user_id)
    
    resume_text = ""
    
    if 'file' in request.files and request.files['file'].filename != '':
        file = request.files['file']
        try:
            if file.mimetype == "application/pdf":
                resume_text = get_pdf_text(file)
            elif file.mimetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                resume_text = get_docx_text(file)
            else:
                return jsonify({"error": "Unsupported file type"}), 400
        except Exception as e:
            return jsonify({"error": f"Error processing file: {str(e)}"}), 500
    
    elif request.json and 'text' in request.json:
        resume_text = request.json.get('text')
    
    if not resume_text:
        return jsonify({"error": "No resume file or text provided."}), 400

    try:
        docs = [Document(page_content=resume_text)]
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = text_splitter.split_documents(docs)

        # Use a per-user chroma directory
        user_db_dir = os.path.join(CHROMA_DB_PATH, user_id)
        os.makedirs(user_db_dir, exist_ok=True)
        try:
            # load existing user store if present
            if os.path.exists(user_db_dir) and os.listdir(user_db_dir):
                user_vs = Chroma(persist_directory=user_db_dir, embedding_function=embeddings)
                try:
                    user_vs.add_documents(chunks)
                    user_vs.persist()
                except Exception as e:
                    print(f"Warning: failed to append to user vector store: {e}")
                    # fallback: recreate
                    user_vs = Chroma.from_documents(documents=chunks, embedding=embeddings, persist_directory=user_db_dir)
                    user_vs.persist()
            else:
                user_vs = Chroma.from_documents(documents=chunks, embedding=embeddings, persist_directory=user_db_dir)
                user_vs.persist()
        except Exception as e:
            print(f"Error handling user vector store: {e}")
            return jsonify({"error": str(e)}), 500

        # deduplicate by hash
        sha = hashlib.sha256(resume_text.encode('utf-8')).hexdigest()
        existing_hashes = {d.get('hash') for d in ingested_docs if d.get('hash')}
        if sha in existing_hashes:
            # Already ingested; return success but indicate duplicate
            return jsonify({"feedback": feedback, "resume_text": resume_text, "note": "duplicate"})

        ingested_docs.append({
            'source': request.files['file'].filename if 'file' in request.files and request.files['file'].filename else 'text-input',
            'text': resume_text,
            'hash': sha,
            'timestamp': datetime.utcnow().isoformat()
        })
        save_ingested_docs(ingested_docs, user_id)

        initial_prompt = f"""
            You are an experienced HR recruiter and career coach.
            Review the following resume text and provide feedback.
            Instructions:
            - Identify exactly 3 key strengths (skills, experiences, or achievements).
            - Identify exactly 3 areas for improvement (clarity, formatting, missing skills, etc).
            - Be concise and use simple language that a fresher can understand.
            - You MUST ONLY respond with a valid JSON object. Do not include any other text, greetings, or explanations.
            Output format:
            {{
                "strengths": ["point 1", "point 2", "point 3"],
                "improvements": ["point 1", "point 2", "point 3"]
            }}
            Resume:
            {resume_text}
        """
        response = chat_model.invoke(initial_prompt).content
        
        json_match = re.search(r'\{.*\}', response, re.DOTALL)
        if json_match:
            json_string = json_match.group(0)
            feedback = json.loads(json_string)
        else:
            raise ValueError("Could not find a valid JSON object in the AI's response.")

        # Include the raw extracted resume text so the frontend can display/store it
        return jsonify({"feedback": feedback, "resume_text": resume_text})

    except Exception as e:
        print(f"Error in process_resume: {e}")
        return jsonify({"error": str(e)}), 500

@app.route("/api/chat", methods=["POST"])
def chat():
    data = request.json
    message = data.get("message")
    if not message:
        return jsonify({"error": "Message is required"}), 400

    # determine user id for per-user vector store
    user_id = get_user_id_from_request(request) or data.get('user_id') or request.args.get('user_id') or 'default'
    user_db_dir = os.path.join(CHROMA_DB_PATH, user_id)
    try:
        if os.path.exists(user_db_dir) and os.listdir(user_db_dir):
            user_vs = Chroma(persist_directory=user_db_dir, embedding_function=embeddings)
        else:
            return jsonify({"error": "Please upload your resume first."}), 400
    except Exception:
        return jsonify({"error": "Please upload your resume first."}), 400

    try:
        prompt_template = ChatPromptTemplate.from_template("""
            You are a helpful and professional resume assistant and career coach.
            Answer the user's question. If the question is about the provided resume, use the context.
            If the question is a general career or skill question, use your broader knowledge.
            
            Context:
            {context}
            
            Question: {input}
        """)

        document_chain = create_stuff_documents_chain(llm=chat_model, prompt=prompt_template)
        retrieval_chain = create_retrieval_chain(retriever=user_vs.as_retriever(), combine_docs_chain=document_chain)
        result = retrieval_chain.invoke({"input": message})
        reply_text = result['answer']

        # --- Skill extraction step: if user message potentially contains skills or self-declared skills,
        # ask the model to extract a short list of skills/keywords and persist them to the user's store.
        try:
            skill_prompt = f"Extract skills or technologies mentioned in this user message as a JSON array of strings. Message: {message}"
            skill_resp = chat_model.invoke(skill_prompt).content
            json_match = re.search(r'\[.*\]', skill_resp, re.DOTALL)
            extracted_skills = []
            if json_match:
                try:
                    extracted_skills = json.loads(json_match.group(0))
                except Exception:
                    extracted_skills = []
            else:
                # fallback regex for common tokens
                possible = re.findall(r"\b(Python|JavaScript|React|Node|SQL|Docker|Kubernetes|AWS|Azure|Java|C#|Git|TypeScript)\b", message, re.IGNORECASE)
                extracted_skills = list({s for s in possible})

            if extracted_skills:
                skill_text = ' '.join(extracted_skills)
                skill_doc = Document(page_content=f"Skills: {skill_text}")
                try:
                    user_vs.add_documents([skill_doc])
                    user_vs.persist()
                except Exception as e:
                    print(f"Warning: failed to add skill doc to user Chroma: {e}")

                ingested = load_ingested_docs(user_id)
                if not any(d.get('text') == skill_text for d in ingested):
                    ingested.append({'source': 'chat-skill', 'text': skill_text, 'timestamp': datetime.utcnow().isoformat()})
                    save_ingested_docs(ingested, user_id)

                reply_text = reply_text + "\n\n(PS: I captured these skills you mentioned: " + ', '.join(extracted_skills) + ")"

        except Exception as skill_e:
            print(f"Skill extraction error: {skill_e}")

        return jsonify({"reply": reply_text})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# --- ENDPOINT 3: Agent Goal Planning (ENHANCED) ---
# --- ENDPOINT 3: Agent Goal Planning ---
@app.route("/api/agent-plan", methods=["POST"])
def agent_plan():
    data = request.json
    goal = data.get("goal")
    
    if not goal:
        return jsonify({"error": "Goal is required"}), 400

    try:
        agent_prompt = f"""
            You are an expert AI agent that helps users create actionable plans to achieve their goals.
            
            Instructions:
            - Take the user's goal and break it down into 10-12 granular, sequential micro-steps the user can follow.
            - For each micro-step provide:
              * step title (short),
              * description (one short sentence),
              * keywords (3-6 keywords or tools to learn or use),
              * exact actions (2-4 very specific tasks the user should do next).
            - The tone should be motivating, concrete, and beginner-friendly.
            - IMPORTANT: The output MUST be strict JSON and you MUST ONLY respond with the JSON object.
            
            Output format:
            {{
                "goal": "{goal}",
                "plan": [
                    {{"step": "title", "description": "short", "keywords": ["k1","k2"], "actions": ["do x","do y"]}},
                    ... (10-12 items)
                ]
            }}

            User's Goal: {goal}
        """

        response = chat_model.invoke(agent_prompt).content
        
        # --- FIX: Two-step robust JSON extraction and cleanup ---
        
        # 1. Regex search for the raw JSON block
        json_match = re.search(r'\{.*\}', response, re.DOTALL)
        
        if json_match:
            json_string = json_match.group(0)
            
            # 2. Add an extra layer of defense: fix common JSON structural errors before parsing
            # This is done by replacing problematic newlines/tabs and escaping quotes inside the string.
            # For simplicity, we'll try to find and repair internal quotes that break the format.
            try:
                plan = json.loads(json_string)
            except json.JSONDecodeError:
                # If standard parse fails, try a manual, safer replacement (often fixes the delimiter issue)
                # This complex replacement logic is best placed in a robust helper, but for this context:
                safe_json_string = json_string.replace('": "', '": "').replace('", "', '", "')
                plan = json.loads(safe_json_string)
            
        else:
            # If no JSON block is found, raise an error
            raise ValueError("Could not find a valid JSON object in the AI's response.")
        
        return jsonify({"plan": plan})
    
    except Exception as e:
        print(f"Error in agent_plan: {e}")
        return jsonify({"error": str(e)}), 500

@app.route("/api/agent-query", methods=["POST"])
def agent_query():
    data = request.json
    query = data.get("query")
    chat_history = data.get("chat_history", [])
    persona = data.get("persona", "a professional career coach")

    if not query:
        return jsonify({"error": "Query is required"}), 400

    try:
        # 🔹 Convert chat history to LangChain-style messages
        history_messages = []
        for msg in chat_history:
            if msg["sender"] == "user":
                history_messages.append(("human", msg["text"]))
            else:
                history_messages.append(("ai", msg["text"]))

        response = agent_executor.invoke({
            "input": query,
            "chat_history": history_messages,
            "persona": persona
        })
        return jsonify({"reply": response.get("output", "No response generated.")})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# --- ENDPOINT 5: Success Prediction Model (NEW FEATURE) ---
@app.route("/api/predict-success", methods=["POST"])
def predict_success():
    data = request.json
    resume_text = data.get("resumeText")
    goal = data.get("goal")

    if not resume_text or not goal:
        return jsonify({"error": "Resume and goal are required."}), 400

    try:
        # NOTE: This prompt tells the AI to act as a prediction model.
        prediction_prompt = f"""
            You are a professional career analyst and data scientist.
            Your task is to analyze the provided resume against the target career goal and predict the likelihood of success.

            Instructions: 
            - Provide a success score as a percentage (from 0 to 100).
            - Write a detailed justification (3-4 sentences) for the score, explaining key strengths and the biggest gaps.
            - You MUST ONLY respond with a valid JSON object.

            Output format: {{ "success_score": 75, "justification": "Based on the resume, the user has strong skills in X and Y... However, there is a gap in Z." }}
            Resume: {resume_text}
            Career Goal: {goal}
        """

        response = chat_model.invoke(prediction_prompt).content
        
        json_match = re.search(r'\{.*\}', response, re.DOTALL)
        if json_match:
            prediction = json.loads(json_match.group(0))
        else:
            raise ValueError("Could not find a valid JSON object in the AI's response.")
            
        return jsonify({"prediction": prediction})

    except Exception as e:
        print(f"Error in predict_success: {e}")
        return jsonify({"error": str(e)}), 500
    
if __name__ == '__main__':
    if not os.path.exists(CHROMA_DB_PATH):
        os.makedirs(CHROMA_DB_PATH)
    app.run(port=5000, debug=True)


# import os
# import json
# import re
# import shutil # New import for deleting the folder
# from flask import Flask, request, jsonify
# from flask_cors import CORS
# from dotenv import load_dotenv
# from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain.docstore.document import Document
# from langchain.chains import create_retrieval_chain
# from langchain.chains.combine_documents.stuff import create_stuff_documents_chain
# from langchain_core.prompts import ChatPromptTemplate
# from langchain_community.vectorstores import Chroma
# import pypdf
# from docx import Document as DocxDocument
# from langchain_chroma import Chroma



# load_dotenv()
# app = Flask(__name__)
# CORS(app)

# vector_store = None
# genai_api_key = os.getenv("GEMINI_API_KEY")

# chat_model = ChatGoogleGenerativeAI(model="gemini-1.5-flash", google_api_key=genai_api_key)
# embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=genai_api_key)

# CHROMA_DB_PATH = "./chroma_db"

# def get_pdf_text(pdf_file):
#     reader = pypdf.PdfReader(pdf_file)
#     text = ""
#     for page in reader.pages:
#         text += page.extract_text()
#     return text

# def get_docx_text(docx_file):
#     doc = DocxDocument(docx_file)
#     text = ""
#     for para in doc.paragraphs:
#         text += para.text + "\n"
#     return text

# @app.route("/api/process-resume", methods=["POST"])
# def process_resume():
#     global vector_store

#     # --- FIX 1: Delete old vector store before creating a new one ---
#     if os.path.exists(CHROMA_DB_PATH):
#         shutil.rmtree(CHROMA_DB_PATH)
#         print("Old vector store deleted. Creating a new one.")
    
#     resume_text = ""
#     if 'file' in request.files and request.files['file'].filename != '':
#         file = request.files['file']
#         try:
#             if file.mimetype == "application/pdf":
#                 resume_text = get_pdf_text(file)
#             elif file.mimetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#                 resume_text = get_docx_text(file)
#             else:
#                 return jsonify({"error": "Unsupported file type"}), 400
#         except Exception as e:
#             return jsonify({"error": f"Error processing file: {str(e)}"}), 500
    
#     elif request.json and 'text' in request.json:
#         resume_text = request.json.get('text')
    
#     if not resume_text:
#         return jsonify({"error": "No resume file or text provided."}), 400

#     try:
#         docs = [Document(page_content=resume_text)]
#         text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
#         chunks = text_splitter.split_documents(docs)

#         vector_store = Chroma.from_documents(
#             documents=chunks,
#             embedding=embeddings,
#             persist_directory=CHROMA_DB_PATH
#         )
#         vector_store.persist()

#         initial_prompt = f"""
#             You are an experienced HR recruiter and career coach.
#             Review the following resume text and provide feedback.
#             Instructions:
#             - Identify exactly 3 key strengths (skills, experiences, or achievements).
#             - Identify exactly 3 areas for improvement (clarity, formatting, missing skills, etc).
#             - Be concise and use simple language that a fresher can understand.
#             - You MUST ONLY respond with a valid JSON object. Do not include any other text, greetings, or explanations.
#             Output format:
#             {{
#                 "strengths": ["point 1", "point 2", "point 3"],
#                 "improvements": ["point 1", "point 2", "point 3"]
#             }}
#             Resume:
#             {resume_text}
#         """
#         response = chat_model.invoke(initial_prompt).content
        
#         json_match = re.search(r'\{.*\}', response, re.DOTALL)
#         if json_match:
#             json_string = json_match.group(0)
#             feedback = json.loads(json_string)
#         else:
#             raise ValueError("Could not find a valid JSON object in the AI's response.")
            
#         return jsonify({"feedback": feedback})
        
#     except Exception as e:
#         print(f"Error in process_resume: {e}")
#         return jsonify({"error": str(e)}), 500

# @app.route("/api/chat", methods=["POST"])
# def chat():
#     global vector_store
#     if vector_store is None:
#         try:
#             vector_store = Chroma(
#                 persist_directory=CHROMA_DB_PATH,
#                 embedding_function=embeddings
#             )
#         except Exception:
#             return jsonify({"error": "Please upload your resume first."}), 400

#     data = request.json
#     message = data.get("message")
#     if not message:
#         return jsonify({"error": "Message is required"}), 400

#     try:
#         # --- FIX 2: Relaxing the prompt for general questions ---
#         prompt_template = ChatPromptTemplate.from_template("""
#             You are a helpful and professional resume assistant and career coach.
#             Answer the user's question. If the question is about the provided resume, use the context.
#             If the question is a general career or skill question, use your broader knowledge.
            
#             Context:
#             {context}
            
#             Question: {input}
#         """)

#         document_chain = create_stuff_documents_chain(llm=chat_model, prompt=prompt_template)
#         retrieval_chain = create_retrieval_chain(retriever=vector_store.as_retriever(), combine_docs_chain=document_chain)
#         result = retrieval_chain.invoke({"input": message})
#         return jsonify({"reply": result['answer']})
#     except Exception as e:
#         return jsonify({"error": str(e)}), 500

# if __name__ == '__main__':
#     if not os.path.exists(CHROMA_DB_PATH):
#         os.makedirs(CHROMA_DB_PATH)
#     app.run(port=5000, debug=True)

# import os
# import json
# import re
# from flask import Flask, request, jsonify
# from flask_cors import CORS
# from dotenv import load_dotenv

# # LangChain and Gemini Imports
# from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain.docstore.document import Document
# from langchain.chains import create_retrieval_chain
# from langchain.chains.combine_documents.stuff import create_stuff_documents_chain
# from langchain_core.prompts import ChatPromptTemplate
# from langchain_community.vectorstores import Chroma

# from langchain.agents import AgentExecutor, create_tool_calling_agent
# from langchain_core.tools import tool
# from langchain_community.tools.tavily_search import TavilySearchResults
# from langchain_core.prompts import MessagesPlaceholder

# # File parsing imports
# import pypdf
# from docx import Document as DocxDocument

# load_dotenv()
# app = Flask(__name__)
# CORS(app)

# vector_store = None
# genai_api_key = os.getenv("GEMINI_API_KEY")

# chat_model = ChatGoogleGenerativeAI(model="gemini-1.5-flash", google_api_key=genai_api_key)
# embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=genai_api_key)

# # Tavily Search Tool Setup
# tavily_api_key = os.getenv("TAVILY_API_KEY")
# search = TavilySearchResults(api_key=tavily_api_key)
# tools = [search]

# # Prompt for the tool-calling agent
# agent_prompt = ChatPromptTemplate.from_messages(
#     [
#         ("system", "You are a helpful AI assistant that can use tools."),
#         ("human", "{input}"),
#         MessagesPlaceholder(variable_name="agent_scratchpad"),
#     ]
# )

# # Create a tool-calling agent
# agent = create_tool_calling_agent(chat_model, tools, agent_prompt)
# agent_executor = AgentExecutor(agent=agent, tools=tools, verbose=True)

# CHROMA_DB_PATH = "./chroma_db"

# def get_pdf_text(pdf_file):
#     reader = pypdf.PdfReader(pdf_file)
#     text = ""
#     for page in reader.pages:
#         text += page.extract_text()
#     return text

# def get_docx_text(docx_file):
#     doc = DocxDocument(docx_file)
#     text = ""
#     for para in doc.paragraphs:
#         text += para.text + "\n"
#     return text

# @app.route("/api/process-resume", methods=["POST"])
# def process_resume():
#     global vector_store

#     resume_text = ""
#     if 'file' in request.files and request.files['file'].filename != '':
#         file = request.files['file']
#         try:
#             if file.mimetype == "application/pdf":
#                 resume_text = get_pdf_text(file)
#             elif file.mimetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#                 resume_text = get_docx_text(file)
#             else:
#                 return jsonify({"error": "Unsupported file type"}), 400
#         except Exception as e:
#             return jsonify({"error": f"Error processing file: {str(e)}"}), 500
    
#     elif request.json and 'text' in request.json:
#         resume_text = request.json.get('text')
    
#     if not resume_text:
#         return jsonify({"error": "No resume file or text provided."}), 400

#     try:
#         docs = [Document(page_content=resume_text)]
#         text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
#         chunks = text_splitter.split_documents(docs)

#         # New: Check if the vector store exists. If so, clear it.
#         if vector_store is not None:
#             vector_store.delete_collection()
#             print("Existing vector store cleared. Adding new data.")
#         else:
#             print("Creating a new vector store.")

#         # Create or update the Chroma vector store
#         vector_store = Chroma.from_documents(
#             documents=chunks,
#             embedding=embeddings,
#             persist_directory=CHROMA_DB_PATH
#         )
#         vector_store.persist()

#         initial_prompt = f"""
#             You are an experienced HR recruiter and career coach.
#             Review the following resume text and provide feedback.
#             Instructions:
#             - Identify exactly 3 key strengths (skills, experiences, or achievements).
#             - Identify exactly 3 areas for improvement (clarity, formatting, missing skills, etc).
#             - Be concise and use simple language that a fresher can understand.
#             - You MUST ONLY respond with a valid JSON object. Do not include any other text, greetings, or explanations.
#             Output format:
#             {{
#                 "strengths": ["point 1", "point 2", "point 3"],
#                 "improvements": ["point 1", "point 2", "point 3"]
#             }}
#             Resume:
#             {resume_text}
#         """
#         response = chat_model.invoke(initial_prompt).content
        
#         json_match = re.search(r'\{.*\}', response, re.DOTALL)
#         if json_match:
#             json_string = json_match.group(0)
#             feedback = json.loads(json_string)
#         else:
#             raise ValueError("Could not find a valid JSON object in the AI's response.")
            
#         return jsonify({"feedback": feedback})
        
#     except Exception as e:
#         print(f"Error in process_resume: {e}")
#         return jsonify({"error": str(e)}), 500

# @app.route("/api/chat", methods=["POST"])
# def chat():
#     global vector_store
#     if vector_store is None:
#         try:
#             vector_store = Chroma(
#                 persist_directory=CHROMA_DB_PATH,
#                 embedding_function=embeddings
#             )
#         except Exception:
#             return jsonify({"error": "Please upload your resume first."}), 400

#     data = request.json
#     message = data.get("message")
#     if not message:
#         return jsonify({"error": "Message is required"}), 400

#     try:
#         prompt_template = ChatPromptTemplate.from_template("""
#             You are a helpful and professional resume assistant and career coach.
#             Answer the user's question. If the question is about the provided resume, use the context.
#             If the question is a general career or skill question, use your broader knowledge.
            
#             Context:
#             {context}
            
#             Question: {input}
#         """)

#         document_chain = create_stuff_documents_chain(llm=chat_model, prompt=prompt_template)
#         retrieval_chain = create_retrieval_chain(retriever=vector_store.as_retriever(), combine_docs_chain=document_chain)
#         result = retrieval_chain.invoke({"input": message})
#         return jsonify({"reply": result['answer']})
#     except Exception as e:
#         return jsonify({"error": str(e)}), 500

        
        
# #---Agent Planning Endpoint---
 
# @app.route("/api/agent-plan", methods=["POST"])
# def agent_plan():
#     data = request.json
#     goal = data.get("goal")

#     if not goal:
#         return jsonify({"error": "Goal is required"}), 400
#     try:
#         agent_prompt = f"""
#             You are an expert AI agent that helps users create actionable plans to achieve their goals.
            
#             Instructions:
#             - Take the user's goal and break it down into 5 clear, sequential steps.
#             - Provide a brief, one-sentence description for each step.
#             - The tone should be motivating and professional.
#             - You MUST ONLY respond with a valid JSON object. Do not include any other text.
            
#             Output format:
#             {{
#                 "goal": "{goal}",
#                 "plan": [
#                     {{"step": "step 1 title", "description": "brief description"}},
#                     {{"step": "step 2 title", "description": "brief description"}},
#                     {{"step": "step 3 title", "description": "brief description"}},
#                     {{"step": "step 4 title", "description": "brief description"}},
#                     {{"step": "step 5 title", "description": "brief description"}}
#                 ]
#             }}

#             User's Goal: {goal}
#         """

#         response = chat_model.invoke(agent_prompt).content

#         json_match = re.search(r'\{.*\}', response, re.DOTALL)
#         if json_match:
#             json_string = json_match.group(0)
#             plan = json.loads(json_string)
#         else:
#             raise ValueError("Could not find a valid JSON object in the AI's response.")
        
#         return jsonify({"plan": plan})

#     except Exception as e:
#         return jsonify({"error": str(e)}), 500


# # Agent query endpoint

# @app.route("/api/agent-query", methods=["POST"])
# def agent_query():
#     data = request.json
#     query = data.get("query")

#     if not query:
#         return jsonify({"error": "Query is required"}), 400

#     try:
#         response = agent_executor.invoke({"input": query})
#         return jsonify({"reply": response['output']})
#     except Exception as e:
#         return jsonify({"error": str(e)}), 500
# if __name__ == '__main__':
#     if not os.path.exists(CHROMA_DB_PATH):
#         os.makedirs(CHROMA_DB_PATH)
#     app.run(port=5000, debug=True)